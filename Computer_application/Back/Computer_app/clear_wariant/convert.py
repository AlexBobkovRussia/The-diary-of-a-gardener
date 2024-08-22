import sys

sys.path.append('')
from docx import Document
import os


# Конвертация файла txt в Word


def convert_txt_to_docx(txt_file_path, docx_file_path):
    txt_file_path = txt_file_path.replace('\\', '/')
    docx_file_path = docx_file_path.replace('\\', '/')
    doc = Document()
    if ((not os.path.exists(f'{txt_file_path}') and not os.path.exists(f'{docx_file_path}')) or
            (not os.path.exists(f'{txt_file_path}') or not os.path.exists(f'{docx_file_path}'))):
        pass
    else:
        with open(txt_file_path, 'r', encoding='utf-8') as txt_file:
            txt = txt_file.read()
        # if not os.path.exists('doc.docx'):
        #     doc = Document()
        #     doc.save('doc.docx')
        # doc = Document('doc.docx')
        doc.add_paragraph(txt)
        doc.save('doc.docx')
    return doc


convert_txt_to_docx('C:/Users/user/PycharmProjects/Семяна/files/Дневник семян.txt',
                    'C:/Users/user/PycharmProjects/Семяна/doc.docx')


# Конвертация файла txt в pdf


def convert_txt_to_pdf(txt_file_path, docx_file_path):
    pass


# Конвертация файла txt в excel


def convert_txt_to_excel(txt_file_path, docx_file_path):
    pass
