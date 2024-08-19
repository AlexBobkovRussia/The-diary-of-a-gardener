import sys

sys.path.append('')
from tkinter import Tk, Button, Label, END, Entry
import os
from datetime import *
from fpdf import FPDF
from spire.doc import FileFormat, Document

window = Tk()
window.geometry('1268x833+1+70')
window.resizable(False, False)
window.title('Дневник семян')
window.config(bg='green')

list_to_collect_information = []
to_done = 0
lbl_of_error = Label(window, text='Заполните правильно все поля!', font='Arial 25', bg='red', fg='white')

file_to_print = open('Дневник семян для печати.txt', 'w', encoding='utf-8')
file_to_print.close()
file_to_print = open('Дневник семян для печати.txt', 'r+', encoding='utf-8')
lenght = file_to_print.readlines()
if len(lenght) == 0:
    file_to_print.write('|-------------------|------------|------------|------------|----------|---------|\n')
    file_to_print.write('| Название растения |    Дата    |    Дата    |    Дата    | Посажено | Собрано |\n')
    file_to_print.write('|                   |   посева   |    сбора   |   записи   | семян    | урожая  |\n')
    file_to_print.write('|                   |    семян   |   урожая   |            |          |         |\n')
    file_to_print.write('|-------------------|------------|------------|------------|----------|---------|\n')
    for i in range(24):
        file_to_print.write('|                   |            |            |            |          |         |\n')
        file_to_print.write('|-------------------|------------|------------|------------|----------|---------|\n')
    file_to_print.close()


def leap_year(year):
    year = int(year)
    if (year % 4 == 0 or year % 400 == 0) and year % 100 != 0:
        return True
    return False


def check_on_right_date(date):
    date_to_work = date.split('.')
    my_dict = {'01': 31, '02': 28, '03': 31, '04': 30, '05': 31, '06': 30, '07': 31, '08': 31, '09': 30, '10': 31,
               '11': 30, '12': 31}
    my_dict2 = {'01': 31, '02': 29, '03': 31, '04': 30, '05': 31, '06': 30, '07': 31, '08': 31, '09': 30, '10': 31,
                '11': 30, '12': 31}
    if len(date_to_work[0]) == 2:
        if len(date_to_work[1]) == 2:
            if len(date_to_work[2]) == 4:
                if leap_year(date_to_work[2]):
                    if (int(my_dict2[date_to_work[1]]) >= int(date_to_work[0]) > 0) and 0 < int(date_to_work[1]) < 13:
                        return True
                else:
                    if (int(my_dict[date_to_work[1]]) >= int(date_to_work[0]) > 0) and 0 < int(date_to_work[1]) < 13:
                        return True
    return False


def check(date1, date2):
    entr_copy = ''
    for i in date1:
        if i == ',':
            entr_copy += '.'
        elif i in '0123456789.':
            entr_copy += i
    date1 = entr_copy
    entr_copy = ''
    for i in date2:
        if i == ',':
            entr_copy += '.'
        elif i in '0123456789.':
            entr_copy += i
    date2 = entr_copy
    if check_on_right_date(date1) and check_on_right_date(date2):
        a = date1.split('.')
        b = date2.split('.')
        if int(a[2]) < int(b[2]):
            return True
        elif int(a[2]) == int(b[2]):
            if int(a[1]) < int(b[1]):
                return True
            elif int(a[1]) == int(b[1]):
                if int(a[0]) <= int(b[0]):
                    return True
                else:
                    return False
    return False


def collecting_information():
    global entr1
    global entr22
    global entr33
    global entr4
    global entr5
    global to_done
    global list_to_collect_information
    global lbl_of_error
    list_to_collect_information2 = []

    def name_of_plants():
        entr = entr1.get()
        if ''.join(entr.split()).isalpha():
            list_to_collect_information2.append(entr)

    name_of_plants()

    if check(entr22.get(), entr33.get()):
        def date_of_grow_plants():
            entr = entr22.get()
            entr_copy = ''
            for i in entr:
                if i == ',':
                    entr_copy += '.'
                elif i in '0123456789.':
                    entr_copy += i
            entr = entr_copy
            info = entr[:entr.find('.')] + entr[entr.find('.') + 1:entr.rfind('.')] + entr[entr.rfind('.') + 1:]
            if str(entr).count('.') == 2 and info.isdigit():
                if len(entr[:entr.find('.')]) == 2 and len(entr[entr.find('.') + 1:entr.rfind('.')]) == 2 and len(
                        entr[entr.rfind('.') + 1:]) == 4:
                    list_to_collect_information2.append([entr])

        def date_of_collecting_a_harvest():
            entr = entr33.get()
            entr_copy = ''
            for i in entr:
                if i == ',':
                    entr_copy += '.'
                elif i in '0123456789.':
                    entr_copy += i
            entr = entr_copy
            info = entr[:entr.find('.')] + entr[entr.find('.') + 1:entr.rfind('.')] + entr[entr.rfind('.') + 1:]
            if str(entr).count('.') == 2 and info.isdigit():
                if len(entr[:entr.find('.')]) == 2 and len(entr[entr.find('.') + 1:entr.rfind('.')]) == 2 and len(
                        entr[entr.rfind('.') + 1:]) == 4:
                    list_to_collect_information2.append([entr])

        date_of_grow_plants()
        date_of_collecting_a_harvest()

    def count_of_grew_plants():
        entr = entr4.get()
        if str(entr).isdigit():
            list_to_collect_information2.append(entr)

    def count_of_assembled_harvest():
        entr = entr5.get()
        if str(entr).isdigit():
            entr += 'кг'
            list_to_collect_information2.append(entr)
        elif str(entr).isalnum():
            if str(entr)[-2:] == 'кг':
                list_to_collect_information2.append(entr)

    list_to_collect_information2.append(['.'.join(str(datetime.now())[:10].split('-')[::-1])])

    count_of_grew_plants()
    count_of_assembled_harvest()
    if to_done == 1:
        lbl_of_error.destroy()
        lbl_of_error = Label(window, text='Заполните правильно все поля!', font='Arial 25', bg='red', fg='white')
    if len(list_to_collect_information2) != 6:
        to_done = 1
        lbl_of_error.grid(row=14, column=1, columnspan=3)
    else:
        list_to_collect_information2 = tuple(list_to_collect_information2)
        entr1.delete(0, END)
        entr22.delete(0, END)
        entr33.delete(0, END)
        entr4.delete(0, END)
        entr5.delete(0, END)
        to_done = 0
        list3 = list_to_collect_information2
        list_to_collect_information2 = []
        list3 = list(list3)

        def str_split(list_to_split, slice, number):
            c = list_to_split[number]
            list_to_split[number] = list_to_split[number].split()
            a = ''
            j = 0
            while c != ''.join(a.split('   ')):
                a += c[slice * (j): slice * (j + 1)]
                a += '   '
                j += 1
            a = a[:-3]
            d = []
            for i in a.split('   '):
                if not (i.isspace()) and i != '':
                    d.append(i.strip())
            list_to_split[number] = d
            return list_to_split

        list3 = str_split(list3, 17, 0)
        list3 = str_split(list3, 8, -2)
        list3 = str_split(list3, 7, -1)
        list3 = tuple(list3)
        if not os.path.exists('Дневник семян.txt'):
            file = open('Дневник семян.txt', 'w', encoding='utf-8')
            file.close()
        file = open('Дневник семян.txt', 'r+', encoding='utf-8')
        f = file.readlines()
        if len(f) == 0:
            my_dict3 = {'01': 'Январь', '02': 'Февраль', '03': 'Март', '04': 'Апрель', '05': 'Май', '06': 'Июнь',
                        '07': 'Июль', '08': 'Август', '09': 'Сентябрь', '10': 'Октябрь', '11': 'Ноябрь',
                        '12': 'Декабрь'}
            file.write('|-------------------|------------|------------|------------|----------|---------|\n')
            file.write('| Название растения |    Дата    |    Дата    |    Дата    | Посажено | Собрано |\n')
            file.write('|                   |   посева   |    сбора   |   записи   | семян    | урожая  |\n')
            file.write('|                   |    семян   |   урожая   |            |          |         |\n')
            file.write('|-------------------|------------|------------|------------|----------|---------|\n')

        to_write = f"{f'| {list3[0][0]} '}{' ' * (20 - len(f'| {list3[0][0]} '))}|"
        to_write += f"{f' {list3[1]} '}|"
        to_write += f"{f' {list3[2]} '}|"
        to_write += f"{f' {list3[3]} '}|"
        to_write += f"{f' {list3[4]} '}{' ' * (13 - len(f'| {list3[5]} '))}|"
        to_write += f"{f' {list3[5]} '}{' ' * (10 - len(f'| {list3[5]} '))}|\n"
        max_len = -1
        for i in list3:
            if len(i) > max_len:
                max_len = len(i)
        print(max_len)
        for i in range(len(list3)):
            while len(list3[i]) < max_len:
                list3[i].append('')

        file = open('Дневник семян.txt', 'r+', encoding='utf-8')
        f = file.readlines()
        next_len_f = 0
        for i in range(0, len(f) * 54, 54):
            if len(f) < i:
                next_len_f = i
                break

        flag = False
        if max_len == 1:
            if f"| {list3[0][0] + ' ' * (17 - len(list3[0][0]))} | {list3[1][0] + ' ' * (10 - len(list3[1][0]))} | {list3[2][0] + ' ' * (10 - len(list3[2][0]))} | {list3[3][0] + ' ' * (10 - len(list3[3][0]))} | {list3[4][0] + ' ' * (8 - len(list3[4][0]))} | {list3[5][0] + ' ' * (7 - len(list3[5][0]))} |" not in f:
                flag = True
            else:
                flag = False
        else:
            f = [i.rstrip() for i in f]
            list4 = f.copy()
            list4 = list4[1:-1]
            list4 = ''.join(list4)
            list4 = list4.split('|-------------------|------------|------------|------------|----------|---------|')
            to_find = ''
            for i in range(max_len):
                to_find += f"| {list3[0][i] + ' ' * (17 - len(list3[0][i]))} | {list3[1][i] + ' ' * (10 - len(list3[1][i]))} | {list3[2][i] + ' ' * (10 - len(list3[2][i]))} | {list3[3][i] + ' ' * (10 - len(list3[3][i]))} | {list3[4][i] + ' ' * (8 - len(list3[4][i]))} | {list3[5][i] + ' ' * (7 - len(list3[5][i]))}|"
            if to_find not in list4:
                flag = True
        if next_len_f - 54 < len(f) + max_len + 1 <= next_len_f and flag:
            for i in range(max_len):
                file.write(
                    f"| {list3[0][i] + ' ' * (17 - len(list3[0][i]))} | {list3[1][i] + ' ' * (10 - len(list3[1][i]))} | {list3[2][i] + ' ' * (10 - len(list3[2][i]))} | {list3[3][i] + ' ' * (10 - len(list3[3][i]))} | {list3[4][i] + ' ' * (8 - len(list3[4][i]))} | {list3[5][i] + ' ' * (8 - len(list3[5][i]))}|\n")
            if len(f) % 54 == 0:
                file.write('|-------------------|------------|------------|------------|----------|---------|\n')
                file.write('| Название растения |    Дата    |    Дата    |    Дата    | Посажено | Собрано |\n')
                file.write('|                   |   посева   |    сбора   |   записи   | семян    | урожая  |\n')
                file.write('|                   |    семян   |   урожая   |            |          |         |\n')
                file.write('|-------------------|------------|------------|------------|----------|---------|\n')
            file.write('|-------------------|------------|------------|------------|----------|---------|\n')
        elif len(f) + max_len + 1 > next_len_f and flag:
            while len(f) + 1 <= next_len_f:
                file.write('                                          \n')
                file.close()
                file = open('Дневник семян.txt', 'r+', encoding='utf-8')
                f = file.readlines()
            if len(f) % 54 == 0:
                file.write('|-------------------|------------|------------|------------|----------|---------|\n')
                file.write('| Название растения |    Дата    |    Дата    |    Дата    | Посажено | Собрано |\n')
                file.write('|                   |   посева   |    сбора   |   записи   | семян    | урожая  |\n')
                file.write('|                   |    семян   |   урожая   |            |          |         |\n')
                file.write('|-------------------|------------|------------|------------|----------|---------|\n')
                for i in range(max_len):
                    file.write(
                        f"| {list3[0][i] + ' ' * (17 - len(list3[0][i]))} | {list3[1][i] + ' ' * (10 - len(list3[1][i]))} | {list3[2][i] + ' ' * (10 - len(list3[2][i]))} | {list3[3][i] + ' ' * (10 - len(list3[3][i]))} | {list3[4][i] + ' ' * (8 - len(list3[4][i]))} | {list3[5][i] + ' ' * (7 - len(list3[5][i]))}|\n")
            file.write('|-------------------|------------|------------|------------|----------|---------|\n')
        file.close()


title_of_window = Label(window, text='Дневник семян', font='Arial 25', bg='green', fg='white')
title_of_window.grid(row=0, column=1, columnspan=3)

lbl1 = Label(window, text='Название растения', font='Arial 15', bg='green', fg='white')
lbl2 = Label(window, text='Дата посева семян', font='Arial 15', bg='green', fg='white')
lbl3 = Label(window, text='Дата сбора урожая', font='Arial 15', bg='green', fg='white')
lbl4 = Label(window, text='Количество засеянных семян', font='Arial 15', bg='green', fg='white')
lbl5 = Label(window, text='Количество собранного урожая', font='Arial 15', bg='green', fg='white')
lbl6 = Label(window,
             text='Чтобы корректно напечатать текст в файле txt: \nоткройте: Блокнот / Файл / Параметры страницы / Поля(мм) ПРАВОЕ=15ММ, ЛЕВОЕ=15 ММ;\nВЕРХНИЙ КОЛОНТИТУЛ = Дневник семян;\n Сохранить',
             font='Arial 20', bg='blue', fg='white')

lbl1.grid(row=6, column=0)
lbl2.grid(row=6, column=1)
lbl3.grid(row=6, column=2)
lbl4.grid(row=6, column=3)
lbl5.grid(row=6, column=4)
lbl6.grid(row=15, column=0, columnspan=5, stick='we')
entr1 = Entry(window, font='Arial 15')
entr22 = Entry(window, font='Arial 15')
entr33 = Entry(window, font='Arial 15')
entr4 = Entry(window, font='Arial 15')
entr5 = Entry(window, font='Arial 15')

entr1.insert(0, 'Укропaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
entr22.insert(0, '12.04.2023')
entr33.insert(0, '13.04.2023')
entr4.insert(0, '17')
entr5.insert(0, '18')

entr1.grid(row=7, column=0, stick='we', padx=5, pady=5)
entr22.grid(row=7, column=1, stick='we', padx=5, pady=5)
entr33.grid(row=7, column=2, stick='we', padx=5, pady=5)
entr4.grid(row=7, column=3, stick='we', padx=5, pady=5)
entr5.grid(row=7, column=4, stick='we', padx=5, pady=5)


def open_window_txt():
    if not os.path.exists("Дневник семян.txt"):
        file = open('Дневник семян.txt', 'w')
        file.close()
    os.system("notepad.exe Дневник семян.txt")


def notes():
    if not os.path.exists("notes.txt"):
        note = open('notes.txt', 'w')
        note.close()
    os.system("notepad.exe notes.txt")


def open_file_txt_to_print():
    if not os.path.exists("Дневник семян для печати.txt"):
        note = open('Дневник семян для печати.txt', 'w')
        note.close()
    os.system("notepad.exe Дневник семян для печати.txt")


def open_window_docx():
    from docx import Document
    if not os.path.exists("Дневник семян.txt"):
        file = open('Дневник семян.txt', 'w')
        file.close()
    if len(open('Дневник семян.txt', 'r').readlines()) > 1:
        file = open('Дневник семян.txt', 'r', encoding='utf-8')
        list1 = file.readlines()
        list2 = list1.copy()
        file.close()
        list2 = list2[1:-1]
        list3 = ''.join(list2).split(
            '|-------------------|------------|------------|------------|----------|---------|\n')
        print(list3)
        for i in list3:
            list3[list3.index(i)] = i.rstrip().split('\n')
        print(list3)
        for i in range(len(list3)):
            for j in range(len(list3[i])):
                list3[i][j] = list3[i][j][1:-1].split('|')
        print(f'list3 = {list3}')

        columns_from_first_to_third = []
        for i in range(6):
            stroka = ''
            for j in range(3):
                stroka += (str(list3[0][j][i]) + '\n')
            columns_from_first_to_third.append(stroka[:-2])
        print(columns_from_first_to_third)

        def filler(column):
            columns_from_third_to_the_end = []
            help = []
            for j in range(1, len(list3)):
                stroka = ''
                help2 = []
                # for i in range(6):
                for k in range(len(list3[j])):
                    stroka += str(list3[j][k][column] + '\n')
                help2.append(stroka.strip())
                help.append(help2)
            columns_from_third_to_the_end.append(help)
            print(f'columns_from_third_to_the_end = {columns_from_third_to_the_end}')
            for i in range(len(columns_from_third_to_the_end)):
                for j in range(len(columns_from_third_to_the_end[i])):
                    table.cell(j + 1, i + column).text = columns_from_third_to_the_end[i][j]
                    print(f'j, i = {j, i}')

        doc = Document()
        table = doc.add_table(rows=len(list3), cols=6)
        table.style = 'Table Grid'

        for i in range(len(columns_from_first_to_third)):
            table.cell(0, i).text = columns_from_first_to_third[i]

        for j in range(6):
            filler(j)

        doc.save('Дневник семян.docx')
    os.startfile('Дневник семян.docx')


# Дорабатывать работу с pdf через другой модуль!!!
def open_window_pdf():
    if not os.path.exists("Дневниксемян.pdf"):
        pdf = FPDF()
        pdf.add_page()
        pdf.output("Дневниксемян.pdf")
    if os.path.exists("Дневник семян.docx"):
        document = Document()

        document.LoadFromFile("Дневник семян.docx")

        document.SaveToFile("Дневниксемян.pdf", FileFormat.PDF)

        document.Close()
    os.system("start Дневниксемян.pdf")


btn1 = Button(window, text='Отправить', font='Arial 20', command=collecting_information)
btn2 = Button(window, text='Открыть файл txt', font='Arial 20', command=open_window_txt)
btn3 = Button(window, text='Открыть файл docx', font='Arial 20', command=open_window_docx)
btn4 = Button(window, text='Открыть файл pdf', font='Arial 20', command=open_window_pdf)
btn6 = Button(window, text='Заметки', font='Arial 20', command=notes)
btn7 = Button(window, text='Открыть файл txt для рукописного заполнения', font='Arial 20',
              command=open_file_txt_to_print)


def to_bind(event):
    collecting_information()


window.bind('<Return>', to_bind)

btn1.grid(row=8, column=1, stick='we', columnspan=3, padx=5, pady=5)
btn2.grid(row=9, column=1, stick='we', columnspan=3, padx=5, pady=5)
btn7.grid(row=10, column=1, stick='we', columnspan=3, padx=5, pady=5)
btn3.grid(row=11, column=1, stick='we', columnspan=3, padx=5, pady=5)
btn4.grid(row=12, column=1, stick='we', columnspan=3, padx=5, pady=5)
btn6.grid(row=13, column=1, stick='we', columnspan=3, padx=5, pady=5)

window.mainloop()
